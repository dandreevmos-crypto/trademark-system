"""Consent letter DOCX generator."""

import os
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models import ConsentLetter, RightsHolder


def format_date_ru(d: date) -> str:
    """Format date in Russian format."""
    return f"{d.day:02d}.{d.month:02d}.{d.year} г."


def format_date_en(d: date) -> str:
    """Format date in English format."""
    months = {
        1: 'January', 2: 'February', 3: 'March', 4: 'April',
        5: 'May', 6: 'June', 7: 'July', 8: 'August',
        9: 'September', 10: 'October', 11: 'November', 12: 'December'
    }
    day = d.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix} of {months[d.month]} {d.year}"


def format_registration_numbers(numbers: List[str]) -> str:
    """Format registration numbers as comma-separated list."""
    return ', '.join(f'№ {n}' for n in numbers)


def _add_russian_part(doc, consent: ConsentLetter, rights_holder: RightsHolder, rh_address_ru: str):
    """Add Russian part of the consent letter."""
    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(rights_holder.name.upper())
    run.bold = True
    run.font.size = Pt(11)

    # Address
    addr_para = doc.add_paragraph()
    addr_para.add_run(rh_address_ru)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(format_date_ru(consent.document_date))

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('СОГЛАСИЕ')
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.add_run('на использование товарных знаков')

    # Main text
    main_para = doc.add_paragraph()
    main_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_text = f'{rights_holder.name.upper()}, зарегистрированное адресу: {rh_address_ru} '
    main_text += f'(далее -«Правообладатель»), обладатель исключительного права на любое изображение '
    main_text += f'товарного знака «{consent.trademark_name}», включая, но не ограничиваясь на товарные знаки, '
    main_text += f'зарегистрированные в Российской Федерации: {format_registration_numbers(consent.registration_numbers)} '
    main_text += f'(далее – «ТЗ»)'
    main_para.add_run(main_text)

    # Signatory
    signatory_para = doc.add_paragraph()
    signatory_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    signatory_text = f'в лице {consent.signatory_position_ru.lower()}а {consent.signatory_name_ru}, '
    signatory_text += 'действующего на основании Устава, предоставляет право использования ТЗ'
    signatory_para.add_run(signatory_text)

    # Recipient
    recipient_para = doc.add_paragraph()
    recipient_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    recipient_text = f'{consent.recipient_name_ru}'
    if consent.recipient_inn:
        recipient_text += f' (ИНН {consent.recipient_inn}'
    recipient_text += f', юридический адрес: {consent.recipient_address_ru})'
    recipient_para.add_run(recipient_text)

    # Purpose
    purpose_para = doc.add_paragraph()
    purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    purpose_text = ''
    if consent.contract_number:
        contract_date_str = format_date_ru(consent.contract_date) if consent.contract_date else ''
        purpose_text += f'в целях исполнения обязательств по договору поставки №{consent.contract_number}'
        if contract_date_str:
            purpose_text += f' от {contract_date_str}'
        purpose_text += ' '
    purpose_text += f'для {consent.usage_purpose_ru}.'
    purpose_para.add_run(purpose_text)

    # Restrictions
    restrict_para = doc.add_paragraph()
    restrict_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    restrict_para.add_run('Использование товарных знаков иными третьими лицами, кроме перечисленных в настоящем Согласии, не допускается.')

    # Validity
    validity_para = doc.add_paragraph()
    validity_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    validity_text = f'Срок действия Согласия: с даты подписания Согласия до {format_date_ru(consent.valid_until)} включительно.'
    validity_para.add_run(validity_text)

    # Termination
    term_para = doc.add_paragraph()
    term_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    term_text = f'По окончании срока либо в случае досрочного отзыва Согласия Правообладателем, '
    term_text += f'{consent.recipient_name_ru} обязан прекратить использование товарных знаков. '
    term_text += 'Возобновление использования товарных знаков возможно лишь при получении нового письменного согласия от Правообладателя.'
    term_para.add_run(term_text)

    # Signature
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.add_run(f'{consent.signatory_position_ru} ')
    sig_para.add_run(consent.signatory_name_ru)

    company_para = doc.add_paragraph()
    company_para.add_run(rights_holder.name.upper())

    stamp_para = doc.add_paragraph()
    stamp_para.add_run('М.П. (подпись)')


def _add_english_part(doc, consent: ConsentLetter, rights_holder: RightsHolder, rh_address_en: str):
    """Add English part of the consent letter."""
    # Header
    header_en = doc.add_paragraph()
    header_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_en = header_en.add_run(rights_holder.name)
    run_en.bold = True

    addr_en = doc.add_paragraph()
    addr_en.add_run(rh_address_en)

    date_en = doc.add_paragraph()
    date_en.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_en.add_run(format_date_en(consent.document_date))

    # Title
    title_en = doc.add_paragraph()
    title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_en_run = title_en.add_run('CONSENT')
    title_en_run.bold = True
    title_en_run.font.size = Pt(14)

    subtitle_en = doc.add_paragraph()
    subtitle_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_en.add_run('to the Trademarks Use')

    # Main text
    main_en = doc.add_paragraph()
    main_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    main_en_text = f'{rights_holder.name}, registered at {rh_address_en}, '
    main_en_text += f'being holder ("Rightholder") of the exclusive right to any image of the trademark "{consent.trademark_name}", '
    main_en_text += f'including but not limited to trademarks registered in the Russian Federation: '
    main_en_text += f'{format_registration_numbers(consent.registration_numbers)} (the "TM"),'
    main_en.add_run(main_en_text)

    # Signatory
    sig_en = doc.add_paragraph()
    sig_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sig_en.add_run(f'represented by {consent.signatory_name_en}, {consent.signatory_position_en}, ')
    sig_en.add_run('acting on the basis of the Articles of Association, grants the right to use the TM to')

    # Recipient
    rec_en = doc.add_paragraph()
    rec_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rec_en_text = f'{consent.recipient_name_en}'
    if consent.recipient_inn:
        rec_en_text += f' (TIN {consent.recipient_inn}'
    rec_en_text += f', legal address: {consent.recipient_address_en})'
    rec_en.add_run(rec_en_text)

    # Purpose
    purp_en = doc.add_paragraph()
    purp_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    purp_en_text = ''
    if consent.contract_number:
        contract_date_en = format_date_en(consent.contract_date) if consent.contract_date else ''
        purp_en_text += f'in order to fulfill obligations under the supply agreement No. {consent.contract_number}'
        if contract_date_en:
            purp_en_text += f' dated {contract_date_en}'
        purp_en_text += ' '
    purp_en_text += f'for {consent.usage_purpose_en}.'
    purp_en.add_run(purp_en_text)

    # Restrictions
    rest_en = doc.add_paragraph()
    rest_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rest_en.add_run('The use of TM by other third parties, other than those listed in this Consent, is not allowed.')

    # Validity
    val_en = doc.add_paragraph()
    val_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    val_en.add_run(f'The Consent period is from the date of signing the Consent until {format_date_en(consent.valid_until)} inclusive.')

    # Termination
    term_en = doc.add_paragraph()
    term_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    term_en_text = f'At the end of the term or in case of early withdrawal of Consent by the Rightholder, '
    term_en_text += f'{consent.recipient_name_en} is obliged to stop using TM. '
    term_en_text += 'The resumption of the use of trademarks is possible only upon receipt of a new written consent from the Rightholder.'
    term_en.add_run(term_en_text)

    # Signature
    doc.add_paragraph()
    sig_en_block = doc.add_paragraph()
    sig_en_block.add_run(f'{consent.signatory_position_en} ')
    sig_en_block.add_run(consent.signatory_name_en)

    company_en = doc.add_paragraph()
    company_en.add_run(rights_holder.name)

    stamp_en = doc.add_paragraph()
    stamp_en.add_run('L.S. (signature)')


def generate_consent_docx(
    consent: ConsentLetter,
    rights_holder: RightsHolder,
) -> BytesIO:
    """Generate consent letter DOCX document.

    Supports document_language: 'ru', 'en', or 'both'.
    Returns BytesIO buffer with the generated document.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    # Get rights holder addresses
    rh_address_ru = ""
    rh_address_en = ""
    if rights_holder.contact_info:
        rh_address_ru = rights_holder.contact_info.get('address_ru', rights_holder.name)
        rh_address_en = rights_holder.contact_info.get('address_en', rights_holder.name)
    else:
        rh_address_ru = rights_holder.name
        rh_address_en = rights_holder.name

    # Get language preference
    language = getattr(consent, 'document_language', 'both') or 'both'

    # Generate based on language
    if language == 'ru':
        _add_russian_part(doc, consent, rights_holder, rh_address_ru)
    elif language == 'en':
        _add_english_part(doc, consent, rights_holder, rh_address_en)
    else:  # both
        _add_russian_part(doc, consent, rights_holder, rh_address_ru)
        doc.add_page_break()
        _add_english_part(doc, consent, rights_holder, rh_address_en)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def save_consent_docx(
    consent: ConsentLetter,
    rights_holder: RightsHolder,
    output_dir: str = "/app/storage/consents",
) -> str:
    """Generate and save consent letter to file.

    Returns the file path.
    """
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    recipient_short = consent.recipient_name_en.replace(' ', '_')[:30]
    date_str = consent.document_date.strftime('%d.%m.%Y')
    filename = f"Authorization_letter_for_{recipient_short}_{date_str}.docx"
    filepath = os.path.join(output_dir, filename)

    buffer = generate_consent_docx(consent, rights_holder)

    with open(filepath, 'wb') as f:
        f.write(buffer.read())

    return filepath
